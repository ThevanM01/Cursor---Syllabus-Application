from flask import Flask, render_template, request, send_file, url_for, redirect
from docxtpl import DocxTemplate
from datetime import datetime
import os
from docx.shared import RGBColor
from docx.oxml import parse_xml

app = Flask(__name__, static_folder='static')###static folder is for the autofill.js file

@app.route('/')
def landing():
    return render_template('landing.html')

@app.route('/admin')
def admin_dashboard():
    return redirect(url_for('index'))  # Temporarily redirect to existing form

@app.route('/educator')
def educator_dashboard():
    return redirect(url_for('index'))  # Temporarily redirect to existing form

@app.route('/student')
def student_dashboard():
    return render_template('student_dashboard.html')

@app.route('/create-syllabus', methods=['GET', 'POST'])
def index():
    general_fields = [
        {
            'columns': [
                {'id': 'course_code', 'label': 'Course Code', 'type': 'text', 'required': True},
                {'id': 'course_name', 'label': 'Course Name', 'type': 'text', 'required': True},
                {'id': 'course_section', 'label': 'Course Section', 'type': 'text', 'required': True}
            ]
        },
        {
            'columns': [
                {'id': 'professor_name', 'label': 'Professor Name', 'type': 'text', 'required': True},
                {'id': 'semester', 'label': 'Semester', 'type': 'select', 'options': ['Summer', 'Fall', 'Winter'], 'required': True},
                {'id': 'year', 'label': 'Year', 'type': 'number', 'min': '2000', 'max': '2100', 'step': '1', 'required': True}
            ]
        },
        {
            'columns': [
                {'id': 'class_location', 'label': 'Class Location', 'type': 'text', 'required': False},
                {'id': 'class_hours', 'label': 'Class Hours', 'type': 'text', 'required': False},
                {'id': 'office_hours', 'label': 'Office Hours', 'type': 'text', 'required': False}
            ]
        },
        {
            'columns': [
                {'id': 'office', 'label': 'Office', 'type': 'text', 'required': False},
                {'id': 'email', 'label': 'Email', 'type': 'email', 'required': True},
                {'id': 'phone', 'label': 'Phone', 'type': 'text', 'required': False}
            ]
        },
        {
            'columns': [
                {'id': 'course_delivery', 'label': 'Course Delivery', 'type': 'text', 'required': False},
                {'id': 'course_exams', 'label': 'Course Exams', 'type': 'text', 'required': False},
                {'id': '', 'label': '', 'type': 'text', 'required': False, 'hidden': True}
            ]
        },
        {
            'columns': [
                {'id': 'prerequisite', 'label': 'Prerequisite(s)', 'type': 'text', 'required': False},
                {'id': 'ta', 'label': 'TA', 'type': 'text', 'required': False},
                {'id': '', 'label': '', 'type': 'text', 'required': False, 'hidden': True}
            ]
        }
    ]
    
    if request.method == 'POST':
        # Get form data
        course_info = {
            'course_name': request.form['course_name'],
            'course_code': request.form['course_code'],
            'semester': request.form['semester'],
            'year': request.form['year'],
            'professor_name': request.form['professor_name'],
            'email': request.form['email'],
            'class_location': request.form['class_location'],
            'class_hours': request.form['class_hours'],
            'course_delivery': request.form['course_delivery'],
            'course_exams': request.form['course_exams'],
            'office': request.form['office'],
            'phone': request.form['phone'],
            'office_hours': request.form['office_hours'],
            'prerequisite': request.form['prerequisite'],
            'ta': request.form['ta'],
            'course_description': request.form.get('course_description', ''),
            'learning_outcomes': request.form.get('learning_outcomes', '').split('\n'),
            'generated_date': datetime.now().strftime("%Y-%m-%d")
        }
        
        # Process timeline data
        timeline_data = []
        timeline_headers = []  # Store column headers
        timeline_rows = []     # Store row data
        
        # Extract all possible timeline keys and sort them
        timeline_keys = set()
        for key in request.form:
            if key.startswith('timeline['):
                parts = key.split('][')
                if len(parts) == 2:
                    timeline_keys.add(parts[1].rstrip(']'))
        
        # Convert to sorted list to maintain consistent order
        timeline_headers = sorted(list(timeline_keys))
        
        # Organize timeline data into rows
        row_index = 0
        timeline_rows = []
        while f'timeline[{row_index}][week]' in request.form:
            row = []
            for header in timeline_headers:
                form_key = f'timeline[{row_index}][{header}]'
                row.append(request.form.get(form_key, ''))
            timeline_rows.append(row)
            row_index += 1
        
        # Store timeline data in a structured format
        course_info['timeline'] = {
            'headers': timeline_headers,
            'rows': timeline_rows
        }
        
        # Load the template
        doc = DocxTemplate('templates/syllabus_template.docx')
        
        # Render the template with our data
        doc.render(course_info)
        
        # Find the timeline placeholder and replace it with a table
        timeline_paragraph = None
        for paragraph in doc.paragraphs:
            if "<<Course Timeline>>" in paragraph.text:
                timeline_paragraph = paragraph
                break
        
        if timeline_paragraph:
            # Create table with headers
            timeline_data = course_info['timeline']
            table = doc.add_table(rows=1, cols=len(timeline_data['headers']))
            table.style = 'Table Grid'
            
            # Style the header row
            header_cells = table.rows[0].cells
            for i, header in enumerate(timeline_data['headers']):
                cell = header_cells[i]
                cell.text = header
                # Set background color to #8F181D
                cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="8F181D"/>'))
                # Set text color to white
                paragraph = cell.paragraphs[0]
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.bold = True
                
            # Add data rows
            for row_data in timeline_data['rows']:
                row_cells = table.add_row().cells
                for i, value in enumerate(row_data):
                    row_cells[i].text = str(value)
            
            # Replace the placeholder paragraph with the table
            timeline_paragraph._element.addnext(table._tbl)
            timeline_paragraph._element.getparent().remove(timeline_paragraph._element)
        
        # Save the generated document
        output_path = f"generated/syllabus_{course_info['course_code']}.docx"
        doc.save(output_path)
        
        # Send the file to the user
        return send_file(output_path, as_attachment=True)
    
    return render_template('form.html', general_fields=general_fields)

if __name__ == '__main__':
    # Create directories if they don't exist
    os.makedirs('templates', exist_ok=True)
    os.makedirs('generated', exist_ok=True)
    app.run(debug=True)